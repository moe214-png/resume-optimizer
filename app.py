from __future__ import annotations

import io
import json
import os
import re
import secrets
import shutil
import threading
import traceback
import uuid
import zipfile
from datetime import datetime
from html import unescape
from pathlib import Path
from xml.etree import ElementTree as ET

import openai
import PyPDF2
from docx import Document
from flask import Flask, redirect, render_template, request, send_from_directory, session, url_for
from werkzeug.exceptions import HTTPException
from werkzeug.utils import secure_filename


app = Flask(__name__)

BASE_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = BASE_DIR / "static"
UPLOAD_DIR = BASE_DIR / "uploads"
ALLOWED_EXTENSIONS = {"pdf", "docx"}
SUPPORT_EXTENSIONS = {"pdf", "docx", "pptx", "txt", "md", "csv", "json", "html", "htm"}
IMAGE_EXTENSIONS = {"png", "jpg", "jpeg", "webp", "bmp", "tif", "tiff"}
JOB_SOURCE_EXTENSIONS = {"pdf", "docx", "txt", "md", "json", "html", "htm"} | IMAGE_EXTENSIONS
MAX_UPLOAD_SIZE = 40 * 1024 * 1024
WORD_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"
OCR_ENGINE = None
APP_VERSION = "2026-05-09-upload-retention"
PDF_FORMAT_WARNING = "你上传的是 PDF 简历。PDF 只能提取文字后重新生成 Word，无法完整保留原简历版式；如需保留格式，请上传原始 DOCX 简历。"
ADMIN_PASSWORD = os.getenv("ADMIN_PASSWORD", "").strip()
UPLOAD_TTL_SECONDS = int(os.getenv("UPLOAD_TTL_SECONDS", str(2 * 60 * 60)))
JOBS: dict[str, dict] = {}
JOBS_LOCK = threading.Lock()

app.config["MAX_CONTENT_LENGTH"] = MAX_UPLOAD_SIZE


def load_local_env() -> None:
    env_path = BASE_DIR / ".env"
    if not env_path.exists():
        return

    for line in env_path.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        os.environ.setdefault(key.strip(), value.strip().strip('"').strip("'"))


load_local_env()
openai.api_base = os.getenv("DEEPSEEK_API_BASE", "https://api.deepseek.com/v1")
openai.api_key = os.getenv("DEEPSEEK_API_KEY", "")
APP_PASSWORD = os.getenv("APP_PASSWORD", "").strip()
SESSION_SECRET = os.getenv("APP_SECRET_KEY", os.getenv("FLASK_SECRET_KEY", "dev-resume-tool-secret"))
OUTPUT_TTL_SECONDS = int(os.getenv("OUTPUT_TTL_SECONDS", str(24 * 60 * 60)))
app.secret_key = SESSION_SECRET


def cleanup_old_outputs() -> None:
    if OUTPUT_TTL_SECONDS <= 0 or not OUTPUT_DIR.exists():
        return

    cutoff = datetime.now().timestamp() - OUTPUT_TTL_SECONDS
    for path in OUTPUT_DIR.glob("optimized_resume_*.docx"):
        try:
            if path.stat().st_mtime < cutoff:
                path.unlink()
        except OSError:
            continue


def cleanup_old_uploads() -> None:
    if UPLOAD_TTL_SECONDS <= 0 or not UPLOAD_DIR.exists():
        return

    cutoff = datetime.now().timestamp() - UPLOAD_TTL_SECONDS
    for path in UPLOAD_DIR.iterdir():
        try:
            if path.is_dir() and path.stat().st_mtime < cutoff:
                shutil.rmtree(path, ignore_errors=True)
        except OSError:
            continue


def password_enabled() -> bool:
    return bool(APP_PASSWORD)


def is_authenticated() -> bool:
    return not password_enabled() or session.get("authenticated") is True


def render_index(**context):
    context.setdefault("password_enabled", password_enabled())
    context.setdefault("is_authenticated", is_authenticated())
    return render_template("index.html", **context)


def set_job_progress(job_id: str, percent: int, message: str, status: str = "running", **extra) -> None:
    with JOBS_LOCK:
        job = JOBS.get(job_id)
        if not job:
            return
        job.update(
            {
                "status": status,
                "percent": max(0, min(100, int(percent))),
                "message": message,
                "updated_at": datetime.now().isoformat(timespec="seconds"),
                **extra,
            }
        )


def get_job(job_id: str) -> dict | None:
    with JOBS_LOCK:
        job = JOBS.get(job_id)
        return dict(job) if job else None


def admin_enabled() -> bool:
    return bool(ADMIN_PASSWORD)


def is_admin_authenticated() -> bool:
    return admin_enabled() and session.get("admin_authenticated") is True


def require_admin():
    if not admin_enabled():
        return render_template("error.html", message="管理功能未启用。请先在 Railway 设置 ADMIN_PASSWORD。"), 404
    if not is_admin_authenticated():
        return redirect(url_for("admin_login", next=request.path))
    return None


def save_upload_file(job_id: str, field: str, filename: str, data: bytes) -> str:
    safe_name = secure_filename(filename) or f"{field}.bin"
    target_dir = UPLOAD_DIR / job_id
    target_dir.mkdir(parents=True, exist_ok=True)
    target = target_dir / f"{field}_{safe_name}"
    target.write_bytes(data)
    return str(target.relative_to(BASE_DIR))


def upload_records() -> list[dict]:
    if not UPLOAD_DIR.exists():
        return []

    records = []
    for job_dir in sorted(UPLOAD_DIR.iterdir(), key=lambda p: p.stat().st_mtime, reverse=True):
        if not job_dir.is_dir():
            continue
        files = []
        for path in sorted(job_dir.iterdir()):
            if path.is_file():
                files.append(
                    {
                        "name": path.name,
                        "size": path.stat().st_size,
                        "download_url": url_for("admin_download_upload", job_id=job_dir.name, filename=path.name),
                    }
                )
        records.append(
            {
                "job_id": job_dir.name,
                "updated_at": datetime.fromtimestamp(job_dir.stat().st_mtime).strftime("%Y-%m-%d %H:%M:%S"),
                "files": files,
                "delete_url": url_for("delete_uploads", job_id=job_dir.name),
            }
        )
    return records


@app.before_request
def run_request_maintenance():
    cleanup_old_outputs()
    cleanup_old_uploads()
    return None


@app.route("/login", methods=["GET", "POST"])
def login():
    if not password_enabled():
        return redirect(url_for("index"))

    error = ""
    if request.method == "POST":
        password = request.form.get("password", "")
        if secrets.compare_digest(password, APP_PASSWORD):
            session["authenticated"] = True
            return redirect(request.args.get("next") or url_for("index"))
        error = "口令不正确，请重试。"
    return render_template("login.html", error=error)


@app.post("/logout")
def logout():
    session.clear()
    return redirect(url_for("index"))


def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def allowed_support_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in SUPPORT_EXTENSIONS


def allowed_job_source_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in JOB_SOURCE_EXTENSIONS


def decode_text_bytes(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "gbk"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def ocr_unavailable_message(filename: str) -> str:
    return (
        f"{filename} 没有可复制文本，可能是图片型 PDF 或截图。"
        "如需自动识别图片文字，请安装 OCR 依赖：PyMuPDF、Pillow、rapidocr-onnxruntime。"
    )


def get_ocr_engine():
    global OCR_ENGINE
    if OCR_ENGINE is not None:
        return OCR_ENGINE
    try:
        from rapidocr_onnxruntime import RapidOCR
    except Exception:
        OCR_ENGINE = False
        return None
    OCR_ENGINE = RapidOCR()
    return OCR_ENGINE


def ocr_image_bytes(data: bytes) -> str:
    engine = get_ocr_engine()
    if not engine:
        return ""

    result, _elapsed = engine(data)
    if not result:
        return ""
    return "\n".join(str(item[1]).strip() for item in result if len(item) >= 2 and str(item[1]).strip())


def ocr_pdf_bytes(data: bytes) -> str:
    try:
        import fitz
    except Exception:
        return ""

    engine = get_ocr_engine()
    if not engine:
        return ""

    texts: list[str] = []
    with fitz.open(stream=data, filetype="pdf") as pdf:
        for page_index, page in enumerate(pdf, start=1):
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            page_text = ocr_image_bytes(pixmap.tobytes("png")).strip()
            if page_text:
                texts.append(f"第 {page_index} 页\n{page_text}")
    return "\n\n".join(texts)


def strip_html(text: str) -> str:
    text = re.sub(r"(?is)<(script|style).*?>.*?</\1>", " ", text)
    text = re.sub(r"(?s)<[^>]+>", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def clean_job_text(text: str) -> str:
    text = unescape(text or "")
    text = re.sub(r"\\u([0-9a-fA-F]{4})", lambda match: chr(int(match.group(1), 16)), text)
    text = re.sub(r"\\[rnt]", "\n", text)
    text = re.sub(r"[ \t\f\v]+", " ", text)
    text = re.sub(r"\n\s*\n+", "\n", text)
    text = re.sub(r"(职位描述|岗位职责|任职要求|职位要求|工作地址|公司介绍)", r"\n\1", text)
    return text.strip()


def flatten_json_values(value, keys: tuple[str, ...], output: list[str]) -> None:
    if isinstance(value, dict):
        for key, item in value.items():
            if key in keys:
                if isinstance(item, list):
                    output.append("、".join(str(part) for part in item if part))
                elif item not in (None, ""):
                    output.append(str(item))
            flatten_json_values(item, keys, output)
    elif isinstance(value, list):
        for item in value:
            flatten_json_values(item, keys, output)


def format_job_json(data: dict) -> str:
    keys = (
        "jobName",
        "positionName",
        "salaryDesc",
        "experienceName",
        "jobExperience",
        "degreeName",
        "jobDegree",
        "locationName",
        "address",
        "brandName",
        "companyName",
        "brandIndustry",
        "industryName",
        "scaleName",
        "brandScaleName",
        "stageName",
        "showSkills",
        "skills",
        "jobLabels",
        "postDescription",
        "jobDescription",
        "description",
        "introduce",
    )
    values: list[str] = []
    flatten_json_values(data, keys, values)

    lines: list[str] = []
    seen: set[str] = set()
    for value in values:
        for line in clean_job_text(value).splitlines():
            line = clean_job_text(strip_html(line))
            if len(line) < 2 or line in seen:
                continue
            seen.add(line)
            lines.append(line)
    return "\n".join(lines)


def focus_job_description_text(text: str) -> str:
    text = clean_job_text(text)
    if not text:
        return ""

    start_patterns = ("职位描述", "岗位职责", "工作职责", "职位要求", "任职要求")
    end_patterns = (
        "竞争力分析",
        "BOSS安全提示",
        "工商信息",
        "工作地址",
        "看过该职位的人还看了",
        "更多职位",
        "精选职位",
        "热门职位",
        "城市招聘",
    )

    start_positions = [text.find(pattern) for pattern in start_patterns if text.find(pattern) != -1]
    if start_positions:
        text = text[min(start_positions) :]

    end_positions = [text.find(pattern) for pattern in end_patterns if text.find(pattern) != -1]
    if end_positions:
        text = text[: min(end_positions)]

    return text.strip()


def extract_job_description_from_bytes(data: bytes, filename: str) -> tuple[str, str | None]:
    if not allowed_job_source_file(filename):
        return "", f"{filename} 暂不支持作为岗位来源，请上传 HTML、TXT、PDF 或 DOCX。"

    text, error = extract_text_from_bytes(data, filename)
    if error:
        return "", error

    ext = filename.rsplit(".", 1)[-1].lower()
    if ext in {"html", "htm"}:
        extracted = extract_job_text_from_html(decode_text_bytes(data))
        if len(extracted) >= 30:
            text = extracted

    focused_text = focus_job_description_text(text)
    return (focused_text or clean_job_text(text))[:12000], None


def extract_job_description_from_file(file) -> tuple[str, str | None]:
    if not file or not file.filename:
        return "", None
    return extract_job_description_from_bytes(file.read(), file.filename)


def extract_job_text_from_html(html: str) -> str:
    snippets: list[str] = []

    for pattern in (
        r'<meta\s+name=["\']description["\']\s+content=["\'](.*?)["\']',
        r'<meta\s+property=["\']og:description["\']\s+content=["\'](.*?)["\']',
        r'<title[^>]*>(.*?)</title>',
    ):
        snippets.extend(re.findall(pattern, html, flags=re.IGNORECASE | re.DOTALL))

    for script_match in re.findall(r'<script[^>]*type=["\']application/ld\+json["\'][^>]*>(.*?)</script>', html, flags=re.IGNORECASE | re.DOTALL):
        try:
            snippets.append(format_job_json(json.loads(unescape(script_match))))
        except Exception:
            snippets.append(script_match)

    for state_match in re.findall(r"__INITIAL_STATE__\s*=\s*(\{.*?\})\s*</script>", html, flags=re.IGNORECASE | re.DOTALL):
        try:
            snippets.append(format_job_json(json.loads(unescape(state_match))))
        except Exception:
            snippets.append(state_match)

    for pattern in (
        r'"(?:jobName|positionName|jobTitle)"\s*:\s*"([^"]+)"',
        r'"(?:salaryDesc|salary|jobSalary)"\s*:\s*"([^"]+)"',
        r'"(?:brandName|companyName)"\s*:\s*"([^"]+)"',
        r'"(?:jobExperience|experienceName)"\s*:\s*"([^"]+)"',
        r'"(?:jobDegree|degreeName)"\s*:\s*"([^"]+)"',
        r'"(?:cityName|businessDistrict)"\s*:\s*"([^"]+)"',
        r'"(?:postDescription|jobDescription|description)"\s*:\s*"([^"]{20,})"',
        r'"(?:skills|jobLabels|labels)"\s*:\s*\[(.*?)\]',
    ):
        snippets.extend(re.findall(pattern, html, flags=re.IGNORECASE | re.DOTALL))

    body_text = strip_html(html)
    if body_text:
        snippets.append(body_text)

    lines: list[str] = []
    seen: set[str] = set()
    for snippet in snippets:
        text = clean_job_text(snippet)
        text = re.sub(r'["{}\[\],:]+', " ", text)
        for line in re.split(r"[\n。；;]+", text):
            line = clean_job_text(line)
            if len(line) < 2 or line in seen:
                continue
            seen.add(line)
            lines.append(line)

    return "\n".join(lines[:120]).strip()


def pptx_slide_texts(data: bytes) -> list[str]:
    texts: list[str] = []
    with zipfile.ZipFile(io.BytesIO(data)) as pptx_zip:
        slide_names = sorted(
            (
                name
                for name in pptx_zip.namelist()
                if re.fullmatch(r"ppt/slides/slide\d+\.xml", name)
            ),
            key=lambda name: int(re.search(r"slide(\d+)\.xml", name).group(1)),
        )
        for index, slide_name in enumerate(slide_names, start=1):
            root = ET.fromstring(pptx_zip.read(slide_name))
            slide_text = "\n".join(
                node.text or ""
                for node in root.iter()
                if node.tag.endswith("}t") and node.text
            ).strip()
            if slide_text:
                texts.append(f"第 {index} 页\n{slide_text}")
    return texts


def extract_text_from_bytes(data: bytes, filename: str) -> tuple[str, str | None]:
    ext = filename.rsplit(".", 1)[-1].lower()

    try:
        if ext == "pdf":
            reader = PyPDF2.PdfReader(io.BytesIO(data))
            text = "\n".join(page.extract_text() or "" for page in reader.pages).strip()
            if not text:
                text = ocr_pdf_bytes(data).strip()
        elif ext == "docx":
            text = "\n".join(docx_paragraph_texts(data)).strip()
        elif ext == "pptx":
            text = "\n\n".join(pptx_slide_texts(data)).strip()
        elif ext in {"txt", "md", "csv", "json"}:
            text = decode_text_bytes(data).strip()
        elif ext in {"html", "htm"}:
            text = strip_html(decode_text_bytes(data))
        elif ext in IMAGE_EXTENSIONS:
            text = ocr_image_bytes(data).strip()
        else:
            return "", f"{filename} 的格式暂不支持。"
    except Exception as exc:
        return "", f"{filename} 解析失败：{exc}"

    if not text and ext in {"pdf"} | IMAGE_EXTENSIONS:
        return "", ocr_unavailable_message(filename)
    if not text:
        return "", f"{filename} 没有读取到有效文本。"
    return text, None


def iter_doc_paragraphs(container, seen_cells: set[int] | None = None):
    if seen_cells is None:
        seen_cells = set()

    for paragraph in container.paragraphs:
        yield paragraph

    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_id = id(cell._tc)
                if cell_id in seen_cells:
                    continue
                seen_cells.add(cell_id)
                yield from iter_doc_paragraphs(cell, seen_cells)


def docx_xml_root(docx_bytes: bytes) -> ET.Element:
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as docx_zip:
        document_xml = docx_zip.read("word/document.xml")
    return ET.fromstring(document_xml)


def docx_paragraph_texts(docx_bytes: bytes) -> list[str]:
    root = docx_xml_root(docx_bytes)
    paragraphs = root.findall(".//w:p", WORD_NS)
    texts = []
    for paragraph in paragraphs:
        text_nodes = paragraph.findall(".//w:t", WORD_NS)
        texts.append("".join(node.text or "" for node in text_nodes))
    return texts


def replace_text_nodes_keep_structure(text_nodes: list[ET.Element], new_text: str) -> None:
    if not text_nodes:
        return

    original_lengths = [len(node.text or "") for node in text_nodes]
    if sum(original_lengths) == 0:
        text_nodes[0].text = new_text
        text_nodes[0].set(XML_SPACE, "preserve")
        for node in text_nodes[1:]:
            node.text = ""
        return

    cursor = 0
    for index, node in enumerate(text_nodes):
        if index == len(text_nodes) - 1:
            node.text = new_text[cursor:]
        else:
            take = original_lengths[index]
            node.text = new_text[cursor : cursor + take]
            cursor += take

        if node.text != (node.text or "").strip():
            node.set(XML_SPACE, "preserve")


def build_optimized_docx_bytes(docx_bytes: bytes, optimized_paragraphs: list[dict]) -> bytes:
    ET.register_namespace("w", WORD_NS["w"])
    root = docx_xml_root(docx_bytes)
    paragraphs = root.findall(".//w:p", WORD_NS)
    optimized_map = {item["paragraph_index"]: item for item in optimized_paragraphs}

    for idx, paragraph in enumerate(paragraphs):
        item = optimized_map.get(idx)
        if not item:
            continue

        text_nodes = paragraph.findall(".//w:t", WORD_NS)
        replace_text_nodes_keep_structure(text_nodes, str(item.get("text", "")))

    new_document_xml = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    output = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as source_zip:
        with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as target_zip:
            for entry in source_zip.infolist():
                data = source_zip.read(entry.filename)
                if entry.filename == "word/document.xml":
                    data = new_document_xml
                target_zip.writestr(entry, data)

    return output.getvalue()


def extract_text_from_file(file, filename: str) -> tuple[str, str | None]:
    ext = filename.rsplit(".", 1)[-1].lower()

    try:
        return extract_text_from_bytes(file.read(), filename)
    except Exception as exc:
        return "", f"文件解析失败：{exc}"


def parse_ai_json(content: str) -> list[dict]:
    content = content.strip()
    content = re.sub(r"^```(?:json)?\s*|\s*```$", "", content, flags=re.IGNORECASE | re.DOTALL)

    try:
        data = json.loads(content)
    except json.JSONDecodeError:
        object_start = content.find("{")
        object_end = content.rfind("}")
        array_start = content.find("[")
        array_end = content.rfind("]")
        snippets = []
        if object_start != -1 and object_end > object_start:
            snippets.append(content[object_start : object_end + 1])
        if array_start != -1 and array_end > array_start:
            snippets.append(content[array_start : array_end + 1])

        data = None
        last_error = None
        for snippet in snippets:
            try:
                data = json.loads(snippet)
                break
            except json.JSONDecodeError as exc:
                last_error = exc
        if data is None:
            raise last_error or ValueError("无法解析 AI 返回的 JSON。")

    if isinstance(data, dict) and isinstance(data.get("items"), list):
        data = data["items"]

    if not isinstance(data, list):
        raise ValueError("AI 返回结果不是 JSON 数组。")

    normalized = []
    for item in data:
        if not isinstance(item, dict):
            continue
        normalized.append(
            {
                "paragraph_index": int(item.get("paragraph_index", len(normalized))),
                "text": str(item.get("text", "")),
                "edits": item.get("edits", []) if isinstance(item.get("edits", []), list) else [],
            }
        )
    return normalized


def call_chat_completion(prompt: str) -> str:
    model = os.getenv("DEEPSEEK_MODEL", "deepseek-chat")
    request_timeout = float(os.getenv("DEEPSEEK_REQUEST_TIMEOUT", "30"))

    try:
        if hasattr(openai, "OpenAI"):
            client = openai.OpenAI(api_key=openai.api_key, base_url=openai.api_base)
            response = client.chat.completions.create(
                model=model,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=4000,
                temperature=0.2,
                response_format={"type": "json_object"},
                timeout=request_timeout,
            )
            return response.choices[0].message.content or ""

        response = openai.ChatCompletion.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=4000,
            temperature=0.2,
            response_format={"type": "json_object"},
            request_timeout=request_timeout,
        )
        return response.choices[0].message["content"]
    except Exception as exc:
        raise RuntimeError(f"AI 请求失败：{exc}") from exc


def local_resume_suggestions(resume_text: str, job_desc: str, paragraphs: list[str] | None = None) -> list[dict]:
    paragraphs = paragraphs if paragraphs is not None else [line.strip() for line in resume_text.splitlines()]
    job_keywords = [
        word
        for word in re.findall(r"[\u4e00-\u9fffA-Za-z0-9+#.]{2,}", job_desc)
        if len(word) >= 2
    ]
    top_keywords = list(dict.fromkeys(job_keywords))[:8]

    suggestions = []
    for idx, paragraph in enumerate(paragraphs):
        if not paragraph:
            suggestions.append({"paragraph_index": idx, "text": "", "edits": []})
            continue

        new_text = paragraph
        edits = []

        if idx == 0 and top_keywords:
            new_text = f"{paragraph} | 目标匹配关键词：{'、'.join(top_keywords[:5])}"
            edits.append(
                {
                    "orig": paragraph,
                    "new": new_text,
                    "reason": "补充岗位关键词，便于快速呈现匹配度。",
                }
            )

        if re.search(r"(负责|参与|进行|协助)", paragraph) and not re.search(r"\d|%|万|次|人|家|项", paragraph):
            new_text = f"{new_text}（建议补充可量化结果，如转化率、增长幅度、项目规模或交付周期。）"
            edits.append(
                {
                    "orig": paragraph,
                    "new": new_text,
                    "reason": "经历描述缺少量化成果，建议补充数据增强说服力。",
                }
            )

        suggestions.append({"paragraph_index": idx, "text": new_text, "edits": edits})

    return suggestions


def is_layout_sensitive_paragraph(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return True
    if len(stripped) <= 12:
        return True
    if len(stripped) <= 24 and not re.search(r"[，。；：,.!?！？]", stripped):
        return True
    if "项目主导人" in stripped:
        return True
    if re.search(r"\d{4}\.\d{1,2}\s*-\s*\d{4}\.\d{1,2}|\d{4}\.\d{1,2}\s*-\s*\d{4}\.\d{1,2}", stripped):
        return True
    if re.search(r"[\w.+-]+@[\w.-]+|\d{7,}", stripped):
        return True
    return False


def normalize_optimizations_for_docx(originals: list[str], optimized_paragraphs: list[dict]) -> list[dict]:
    optimized_map = {item.get("paragraph_index"): item for item in optimized_paragraphs if isinstance(item, dict)}
    original_lookup = {text.strip(): idx for idx, text in enumerate(originals) if text.strip()}
    normalized = []

    for idx, original in enumerate(originals):
        item = optimized_map.get(idx, {})
        new_text = str(item.get("text", original))
        edits = item.get("edits", [])
        original_clean = original.strip()

        if is_layout_sensitive_paragraph(original_clean):
            new_text = original
            edits = []
        elif new_text.strip() == original_clean and new_text != original:
            new_text = original
            edits = []
        elif not new_text:
            new_text = original
            edits = []
        else:
            max_len = max(len(original_clean) + 20, int(len(original_clean) * 1.30))
            min_len = int(len(original_clean) * (0.68 if len(original_clean) >= 80 else 0.5))
            duplicates_other_paragraph = original_lookup.get(new_text.strip()) not in (None, idx)
            if len(new_text) > max_len or len(new_text) < min_len or duplicates_other_paragraph:
                new_text = original
                edits = []

        normalized.append({"paragraph_index": idx, "text": new_text, "edits": edits})

    return normalized


def summarize_format_protection(originals: list[str], before_items: list[dict], after_items: list[dict]) -> dict:
    before_map = {item.get("paragraph_index"): item for item in before_items if isinstance(item, dict)}
    protected = []
    for item in after_items:
        if not isinstance(item, dict):
            continue
        idx = item.get("paragraph_index")
        if not isinstance(idx, int) or idx >= len(originals):
            continue
        before_item = before_map.get(idx)
        if not before_item:
            continue
        before_text = str(before_item.get("text", ""))
        after_text = str(item.get("text", ""))
        original_text = originals[idx]
        if before_text.strip() != original_text.strip() and after_text == original_text:
            protected.append(idx)

    return {
        "count": len(protected),
        "paragraphs": protected[:8],
    }


def build_optimization_diagnosis(
    originals: list[str],
    optimized_paragraphs: list[dict],
    changes: list[dict],
    match_analysis: dict,
    format_protection: dict | None = None,
) -> dict:
    total_paragraphs = len([text for text in originals if str(text).strip()])
    candidate_count = len(select_optimizable_paragraphs(originals))
    changed_count = len(changes)
    score = int(match_analysis.get("score", 0) or 0)
    gaps = match_analysis.get("gaps", [])
    suggestions = match_analysis.get("suggestions", [])
    mismatch_details = match_analysis.get("mismatch_details", [])
    format_count = int((format_protection or {}).get("count", 0) or 0)

    reasons: list[str] = []
    next_steps: list[str] = []

    if changed_count == 0:
        reasons.append("本次没有生成实质改写。")
    elif changed_count <= 2:
        reasons.append(f"本次只改写了 {changed_count} 个段落，改动较少。")
    else:
        reasons.append(f"本次改写了 {changed_count} 个段落，主要集中在可与岗位建立证据关系的经历。")

    if score < 55:
        reasons.append("岗位与现有简历匹配度偏低，系统不宜为了贴合岗位而编造不存在的行业、地区、年限、语言或项目经历。")
        if gaps:
            reasons.append("主要不匹配点：" + "；".join(str(item) for item in gaps[:4]))
        if mismatch_details and isinstance(mismatch_details[0], dict) and mismatch_details[0].get("action"):
            next_steps.append(str(mismatch_details[0]["action"]))
        else:
            next_steps.append("建议补充真实经历材料，例如相近项目、目标用户、关键指标、协作范围或业务复盘，再重新生成。")
    elif score < 70:
        reasons.append("岗位与简历有部分交集，但仍存在明显短板，因此只会优先改写能被原经历支撑的段落。")
        if gaps:
            reasons.append("需要重点补足：" + "；".join(str(item) for item in gaps[:4]))
        if mismatch_details and isinstance(mismatch_details[0], dict) and mismatch_details[0].get("action"):
            next_steps.append(str(mismatch_details[0]["action"]))
        else:
            next_steps.append("可以在“过往工作成果补充”里上传项目复盘、周报、数据记录，让系统找到更多可改写证据。")

    if candidate_count == 0:
        reasons.append("简历中可安全改写的正文段落较少，很多内容像姓名、时间、联系方式、短标题或版式敏感字段。")
        next_steps.append("建议上传包含完整工作经历正文的 DOCX，或在简历中增加项目背景、动作、结果三类信息。")
    elif candidate_count <= 3:
        reasons.append(f"系统只识别到 {candidate_count} 个适合改写的正文段落，所以改动空间有限。")

    if format_count:
        reasons.append(f"另有 {format_count} 个 AI 建议被格式保护规则拦下，原因通常是改写过长、过短、重复，或会破坏 Word 原有版式。")
        next_steps.append("如希望放开版式限制，可以改用粘贴文本生成新 Word，或先把简历正文整理成更常规的段落结构。")

    if not next_steps and suggestions:
        next_steps.extend(str(item) for item in suggestions[:3])

    if not reasons:
        reasons.append("系统已按岗位匹配度和简历格式约束完成改写。")

    return {
        "changed_count": changed_count,
        "candidate_count": candidate_count,
        "total_paragraphs": total_paragraphs,
        "match_score": score,
        "format_protected_count": format_count,
        "reasons": reasons,
        "next_steps": next_steps[:4],
    }


def format_numbered_resume(paragraphs: list[str]) -> str:
    return "\n".join(f"[{idx}] {text}" for idx, text in enumerate(paragraphs))


def select_optimizable_paragraphs(paragraphs: list[str], limit: int = 18) -> list[tuple[int, str]]:
    candidates = []
    preferred_keywords = (
        "负责",
        "主导",
        "搭建",
        "运营",
        "策略",
        "数据",
        "增长",
        "孵化",
        "社群",
        "作者",
        "达人",
        "主播",
        "内容",
        "项目",
        "协同",
        "转化",
        "提升",
    )

    for idx, text in enumerate(paragraphs):
        stripped = text.strip()
        if is_layout_sensitive_paragraph(stripped):
            continue
        if len(stripped) < 35:
            continue

        score = sum(1 for keyword in preferred_keywords if keyword in stripped)
        if re.search(r"\d|%|万|\+", stripped):
            score += 2
        if len(stripped) >= 80:
            score += 1
        if score > 0:
            candidates.append((score, idx, text))

    candidates.sort(key=lambda item: (-item[0], item[1]))
    selected = sorted((idx, text) for _score, idx, text in candidates[:limit])
    return selected


def format_candidate_resume(paragraphs: list[str]) -> str:
    candidates = select_optimizable_paragraphs(paragraphs)
    return "\n".join(f"[{idx} | 原文字数 {len(text.strip())}] {text}" for idx, text in candidates)


def format_achievement_analysis_for_prompt(achievement_analysis: dict | None) -> str:
    if not achievement_analysis:
        return "无"

    items = achievement_analysis.get("items", [])
    if not isinstance(items, list) or not items:
        return "无"

    lines = []
    for idx, item in enumerate(items[:8], start=1):
        if not isinstance(item, dict):
            continue
        evidence = str(item.get("evidence", "")).strip()
        relevance = str(item.get("relevance", "")).strip()
        suggested_use = str(item.get("suggested_use", "")).strip()
        confidence = str(item.get("confidence", "")).strip()
        if evidence:
            lines.append(
                f"{idx}. 证据：{evidence}\n"
                f"   匹配原因：{relevance or '与岗位相关'}\n"
                f"   建议用法：{suggested_use or '补充到相关经历'}\n"
                f"   可信度：{confidence or '中'}"
            )
    return "\n".join(lines) if lines else "无"


def local_job_profile(job_desc: str, target_role: str = "") -> dict:
    keywords = list(
        dict.fromkeys(
            word
            for word in re.findall(r"[\u4e00-\u9fffA-Za-z0-9+#.]{2,}", job_desc)
            if len(word) >= 2
        )
    )[:12]
    is_creator_role = any(word in job_desc for word in ("创作者", "达人", "主播", "作者", "UGC", "内容生态"))
    is_community_role = any(word in job_desc for word in ("社区", "UGC", "图文", "短视频", "话题"))
    is_game_role = any(word in job_desc for word in ("游戏", "手游", "休闲", "发行", "LTV", "留存"))
    is_overseas_role = any(word in job_desc for word in ("海外", "欧美", "英语", "本地化"))

    if is_creator_role or is_community_role:
        business_problem = "岗位核心不是简单发内容，而是扩大优质创作者/UGC供给，并通过分层激励、活动话题、内容机制和数据复盘提升留存、活跃与内容质量。"
        ideal_candidate = "理想候选人应做过创作者/达人/作者生态运营，能证明自己负责过拉新、分层、激励、活动策划、内容供给增长和跨团队落地。"
        core_tasks = ["创作者拉新与孵化", "分层运营与激励体系", "内容供给和话题活动", "数据复盘与策略迭代", "跨团队资源协同"]
        must_haves = ["创作者运营方法论", "活动策划能力", "数据分析能力", "内容生态理解", "项目 owner 能力"]
        transferable = ["主播/达人运营", "生命周期管理", "分层激励", "活动增长", "数据复盘", "跨部门协同"]
        evidence = ["创作者规模或增长数据", "留存/活跃/投稿/互动指标", "活动机制和曝光结果", "分层策略和权益设计", "产品/市场/审核协同案例"]
        strategy = "把直播/主播经验翻译成“内容创作者生态运营”语言，突出创作者生命周期、分层激励、内容供给增长和数据复盘；平台或体裁差异要用可迁移案例解释。"
    elif is_game_role:
        business_problem = "岗位核心是围绕游戏用户生命周期和商业化目标，提升留存、活跃、付费和长期生态健康。"
        ideal_candidate = "理想候选人应具备游戏运营、活动设计、用户分层、数据复盘和跨团队推动经验。"
        core_tasks = ["长线运营", "活动设计", "用户分层", "商业化优化", "数据复盘"]
        must_haves = ["游戏运营经验", "数据分析", "活动策划", "用户理解", "协同推进"]
        transferable = ["活动运营", "用户分层", "内容/社群运营", "数据复盘", "项目管理"]
        evidence = ["留存/活跃/LTV/付费指标", "活动从0到1案例", "用户分层策略", "本地化或品类理解", "跨团队落地结果"]
        strategy = "围绕游戏运营核心指标重排经历，突出活动机制、用户分层、商业化或留存提升结果。"
    else:
        business_problem = "招聘方希望候选人能接手岗位核心业务问题，并用过往项目证据证明可迁移能力。"
        ideal_candidate = "具备相近业务场景、完整项目 owner 经历、数据复盘能力和跨团队推进经验的候选人。"
        core_tasks = keywords[:5]
        must_haves = keywords[5:10]
        transferable = ["项目 owner 能力", "用户/创作者分层运营", "活动策划", "数据复盘", "跨团队协作"]
        evidence = ["项目背景", "负责动作", "协作对象", "量化结果", "复盘方法"]
        strategy = "围绕岗位背后的业务问题重排表达重心，用候选人已有经历和成果证明匹配度。"

    hard_gates = []
    if is_overseas_role:
        hard_gates.append("海外/欧美用户理解或英语工作语言证据")

    return {
        "role_summary": f"{target_role or '目标岗位'}需要候选人证明自己能解决该岗位的核心业务问题，而不是只覆盖表层关键词。",
        "business_problem": business_problem,
        "ideal_candidate": ideal_candidate,
        "core_tasks": core_tasks,
        "must_haves": must_haves,
        "evidence_requirements": evidence,
        "transferable_abilities": transferable,
        "hard_gates": hard_gates,
        "hidden_criteria": ["表达要自然可信", "优先呈现可验证成果", "避免堆砌岗位关键词"],
        "resume_strategy": strategy,
        "keywords": keywords,
        "used_ai": False,
        "analysis_error": "未调用 AI，使用本地岗位关键词规则。",
    }


def analyze_job_profile(job_desc: str, target_role: str = "") -> dict:
    if not openai.api_key:
        return local_job_profile(job_desc, target_role)

    prompt = f"""
你是资深招聘顾问。请先深度理解岗位，不要评估简历，只分析 JD 本身。

只返回 JSON 对象，不要 Markdown，格式：
{{
  "role_summary": "一句话说明这个岗位真正要找什么人",
  "business_problem": "这个岗位被招聘进来要解决的业务问题",
  "ideal_candidate": "理想候选人的画像，要具体",
  "core_tasks": ["核心任务1", "核心任务2", "核心任务3"],
  "must_haves": ["硬性能力1", "硬性能力2", "硬性能力3"],
  "evidence_requirements": ["简历里必须看到的证据1", "证据2", "证据3"],
  "transferable_abilities": ["可以从相邻经历迁移过来的能力1", "能力2", "能力3"],
  "hard_gates": ["不能靠话术弥补的硬门槛1", "硬门槛2"],
  "hidden_criteria": ["隐性筛选点1", "隐性筛选点2", "隐性筛选点3"],
  "resume_strategy": "简历应该如何改写才自然、有说服力",
  "keywords": ["关键词1", "关键词2", "关键词3"]
}}

分析要求：
1. 不要只复述 JD，要判断招聘方真正重视的业务问题、能力和证据。
2. 区分“表层关键词”和“背后的能力要求”，尤其要识别哪些能力可以从相邻经验迁移。
3. hard_gates 只写真正不能靠简历表达弥补的条件，例如明确年限、语言、地区/海外经验、特定品类经验。
4. evidence_requirements 要写成招聘方希望在简历中看到的证据形态，例如“分层策略+指标变化”，不要只写关键词。
5. 给出的简历策略要避免生硬堆词，强调怎样把已有经历自然改成岗位语言。

目标岗位：{target_role or "未指定"}

岗位描述：
{job_desc}
""".strip()

    try:
        data = parse_json_object(call_chat_completion(prompt))
        profile = local_job_profile(job_desc, target_role)
        for key in ("role_summary", "resume_strategy", "business_problem", "ideal_candidate"):
            if data.get(key):
                profile[key] = str(data[key]).strip()
        for key in (
            "core_tasks",
            "must_haves",
            "hidden_criteria",
            "keywords",
            "evidence_requirements",
            "transferable_abilities",
            "hard_gates",
        ):
            value = data.get(key, [])
            if isinstance(value, str):
                profile[key] = [value]
            elif isinstance(value, list):
                profile[key] = [str(item).strip() for item in value if str(item).strip()][:8]
        profile["used_ai"] = True
        profile["analysis_error"] = ""
        return profile
    except Exception as exc:
        profile = local_job_profile(job_desc, target_role)
        profile["analysis_error"] = f"岗位深读 AI 调用失败，已退回本地规则：{exc}"
        return profile


def format_job_profile_for_prompt(job_profile: dict | None) -> str:
    if not job_profile:
        return "无"

    def join_items(key: str) -> str:
        value = job_profile.get(key, [])
        if isinstance(value, list):
            return "、".join(str(item) for item in value if str(item).strip()) or "无"
        return str(value) or "无"

    return (
        f"岗位本质：{job_profile.get('role_summary', '无')}\n"
        f"业务问题：{job_profile.get('business_problem', '无')}\n"
        f"理想候选人：{job_profile.get('ideal_candidate', '无')}\n"
        f"核心任务：{join_items('core_tasks')}\n"
        f"硬性能力：{join_items('must_haves')}\n"
        f"证据要求：{join_items('evidence_requirements')}\n"
        f"可迁移能力：{join_items('transferable_abilities')}\n"
        f"硬门槛：{join_items('hard_gates')}\n"
        f"隐性筛选点：{join_items('hidden_criteria')}\n"
        f"关键词：{join_items('keywords')}\n"
        f"简历策略：{job_profile.get('resume_strategy', '无')}"
    )


def build_force_rewrite_prompt(
    paragraphs: list[str],
    job_desc: str,
    target_role: str,
    style_text: str,
    achievement_analysis: dict | None = None,
    job_profile: dict | None = None,
    limit: int = 12,
) -> str:
    candidates = select_optimizable_paragraphs(paragraphs, limit=limit)
    candidate_text = "\n".join(f"[{idx} | 原文字数 {len(text.strip())}] {text}" for idx, text in candidates)
    achievement_text = format_achievement_analysis_for_prompt(achievement_analysis)
    job_profile_text = format_job_profile_for_prompt(job_profile)

    return f"""
你是中文简历改写助手。请只改写下面列出的候选段落，目标是更匹配岗位，但必须保持 Word 版式稳定。

硬性要求：
1. 只返回 JSON，不要输出任何额外文字，格式必须是 {{"items":[...]}}。
2. items 中至少返回 5 个发生实质改写的段落，最多返回 10 个。
3. paragraph_index 必须使用候选段落方括号里的数字。
4. text 最多只能比该段“原文字数”长 30%，允许为了岗位匹配适当增加关键词和表达。
5. 不新增事实、不编造数据、不改公司/学校/日期/联系方式。
6. 不要只改大小写、标点或空格。
7. edits 中必须写清楚 orig、new、reason。

改写方向：
- 根据“岗位理解”改写，不要机械堆关键词。
- 把泛泛的“负责/参与”改得更结果导向、更贴近岗位，但语气要像真实简历。
- 保留原有数字、项目、成果。
- 如果“可补充成果”里有与候选段落相关的信息，可以把它自然补进对应段落。
- 避免反复在句尾追加“契合岗位/具备能力/提升匹配度”这类空泛总结。

目标岗位：{target_role or "未指定"}
优化风格：{style_text}

岗位描述：
{job_desc}

岗位理解：
{job_profile_text}

候选段落：
{candidate_text}

可补充成果：
{achievement_text}
""".strip()


def optimize_resume(
    resume_text: str,
    job_desc: str,
    target_role: str = "",
    style: str = "professional",
    paragraphs: list[str] | None = None,
    achievement_analysis: dict | None = None,
    job_profile: dict | None = None,
) -> list[dict]:
    if not openai.api_key:
        return local_resume_suggestions(resume_text, job_desc, paragraphs)

    style_map = {
        "professional": "专业、清晰、结果导向",
        "concise": "简洁、克制、突出重点",
        "impact": "更强调业务影响、量化结果和岗位关键词",
    }
    style_text = style_map.get(style, style_map["professional"])
    resume_for_prompt = format_candidate_resume(paragraphs) if paragraphs is not None else resume_text
    achievement_text = format_achievement_analysis_for_prompt(achievement_analysis)
    job_profile_text = format_job_profile_for_prompt(job_profile)
    index_instruction = (
        "下面只列出了允许优化的正文段落，格式为 [段落索引 | 原文字数 N]。paragraph_index 必须严格使用方括号里的段落索引。"
        if paragraphs is not None
        else "paragraph_index 对应原简历段落顺序，从 0 开始。"
    )

    prompt = f"""
你是一名资深中文简历优化顾问。请根据岗位描述优化简历文本，使其更贴合目标岗位，但不要编造不存在的经历、公司、学历、证书或数据。

输出要求：
1. 只返回 JSON，不要输出 Markdown 或额外说明，格式必须是 {{"items":[...]}}。
2. items 中每个元素对应一个发生实质优化的原简历段落，字段为：
   - paragraph_index：原段落索引，从 0 开始
   - text：优化后的段落文本
   - edits：修改列表，每项包含 orig、new、reason
3. 保留原简历结构和段落顺序；不需要返回未修改段落。
4. 严格保持原 Word 版式：不要新增段落，不要合并段落，不要添加换行符。
5. 每个段落优化后的字数最多比原文字数 N 长 30%，允许为了岗位匹配适当增加关键词和表达，但不要扩写成新段落。
6. 不要修改姓名、联系方式、邮箱、日期、公司名称、学校名称、章节标题和项目标题。
7. 优先优化“核心优势”“工作经历”“核心项目经历”中的正文段落，必须参考“岗位理解”，不要机械堆关键词。
8. 如果缺少可量化信息，只能提示“建议补充”，不能虚构数字。
9. 不要只做大小写、标点或空格调整；每条修改都要有岗位证据、表达重心或结果导向上的实质提升。
10. 请从给出的允许优化段落中尽量优化 8-12 个正文段落。可以改写句式、替换关键词、调整表达重心，但必须保留原事实和数字。
11. 如果“可补充成果”中有与某段经历直接相关的信息，可以自然合并进该段；但只能使用材料中出现过的真实事实。
12. 改写后的句子要像候选人原本经历的自然升级版，避免在句尾生硬追加“契合岗位/具备能力/提升匹配度”等空泛总结。
13. {index_instruction}

目标岗位：{target_role or "未指定"}
优化风格：{style_text}

岗位描述：
{job_desc}

岗位理解：
{job_profile_text}

允许优化的简历段落：
{resume_for_prompt}

可补充成果：
{achievement_text}
""".strip()

    try:
        content = call_chat_completion(prompt)
        data = parse_ai_json(content)
        changed_count = 0
        if paragraphs is not None:
            originals = {idx: text for idx, text in enumerate(paragraphs)}
            changed_count = sum(
                1
                for item in data
                if isinstance(item, dict)
                and item.get("paragraph_index") in originals
                and str(item.get("text", "")) != originals[item.get("paragraph_index")]
            )

        if paragraphs is not None and changed_count < 3:
            force_prompt = build_force_rewrite_prompt(
                paragraphs,
                job_desc,
                target_role,
                style_text,
                achievement_analysis,
                job_profile,
            )
            return parse_ai_json(call_chat_completion(force_prompt))

        return data
    except Exception as exc:
        fallback = local_resume_suggestions(resume_text, job_desc, paragraphs)
        if fallback:
            fallback[0]["edits"].append(
                {
                    "orig": "",
                    "new": "",
                    "reason": f"AI 优化暂不可用，已生成本地建议：{exc}",
                }
            )
        return fallback


def replace_paragraph_text_keep_format(paragraph, new_text: str) -> None:
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(new_text)
        return

    original_lengths = [len(run.text or "") for run in runs]
    if sum(original_lengths) == 0:
        runs[0].text = new_text
        for run in runs[1:]:
            run.text = ""
        return

    cursor = 0
    for index, run in enumerate(runs):
        if index == len(runs) - 1:
            run.text = new_text[cursor:]
            break

        take = original_lengths[index]
        run.text = new_text[cursor : cursor + take]
        cursor += take


def apply_optimizations_to_doc(doc: Document, optimized_paragraphs: list[dict]) -> Document:
    paragraphs = list(iter_doc_paragraphs(doc))
    original_paragraphs = [p.text for p in paragraphs]
    optimized_map = {item["paragraph_index"]: item for item in optimized_paragraphs}

    for idx, paragraph in enumerate(paragraphs):
        original_text = original_paragraphs[idx] if idx < len(original_paragraphs) else ""
        item = optimized_map.get(idx, {"text": original_text, "edits": []})
        new_text = item.get("text", original_text)
        replace_paragraph_text_keep_format(paragraph, new_text)

    return doc


def build_docx_from_text(optimized_paragraphs: list[dict]) -> Document:
    doc = Document()
    for item in optimized_paragraphs:
        paragraph = doc.add_paragraph()
        paragraph.add_run(item.get("text", ""))
        reasons = [
            str(edit.get("reason", "")).strip()
            for edit in item.get("edits", [])
            if isinstance(edit, dict) and str(edit.get("reason", "")).strip()
        ]
        if reasons:
            note = paragraph.add_run(f"（修改说明：{'；'.join(reasons)}）")
            note.italic = True
    return doc


def build_change_log(originals: list[str], optimized_paragraphs: list[dict]) -> list[dict]:
    changes = []

    for item in optimized_paragraphs:
        idx = item.get("paragraph_index")
        if not isinstance(idx, int) or idx >= len(originals):
            continue

        before = originals[idx]
        after = str(item.get("text", before))
        if before == after:
            continue

        reasons = [
            str(edit.get("reason", "")).strip()
            for edit in item.get("edits", [])
            if isinstance(edit, dict) and str(edit.get("reason", "")).strip()
        ]
        reason = "；".join(dict.fromkeys(reasons)) or "根据岗位描述优化表达，使经历更贴合目标岗位。"

        changes.append(
            {
                "paragraph_index": idx,
                "before": before,
                "after": after,
                "before_len": len(before.strip()),
                "after_len": len(after.strip()),
                "reason": reason,
            }
        )

    return changes


def parse_json_object(content: str) -> dict:
    content = content.strip()
    content = re.sub(r"^```(?:json)?\s*|\s*```$", "", content, flags=re.IGNORECASE | re.DOTALL)

    try:
        data = json.loads(content)
    except json.JSONDecodeError:
        start = content.find("{")
        end = content.rfind("}")
        if start == -1 or end == -1 or end <= start:
            raise
        data = json.loads(content[start : end + 1])

    if not isinstance(data, dict):
        raise ValueError("AI 返回结果不是 JSON 对象。")
    return data


def local_match_analysis(resume_text: str, job_desc: str) -> dict:
    resume_words = set(re.findall(r"[\u4e00-\u9fffA-Za-z0-9+#.]{2,}", resume_text))
    job_words = [
        word
        for word in re.findall(r"[\u4e00-\u9fffA-Za-z0-9+#.]{2,}", job_desc)
        if len(word) >= 2
    ]
    unique_job_words = list(dict.fromkeys(job_words))
    matched = [word for word in unique_job_words if word in resume_words]
    keyword_score = int((len(matched) / max(len(unique_job_words), 1)) * 30)
    transferable_groups = {
        "创作者/达人运营": ("创作者", "作者", "达人", "主播", "签约", "孵化"),
        "生命周期/分层运营": ("生命周期", "分层", "成长", "留存", "活跃", "激励"),
        "活动/内容策划": ("活动", "内容", "策划", "热点", "话题", "曝光"),
        "数据驱动": ("数据", "转化", "留存率", "点击", "ROI", "LTV", "复盘"),
        "跨团队协作": ("协同", "跨部门", "产品", "市场", "赛事", "资源"),
        "社区/UGC迁移": ("社区", "UGC", "短视频", "图文", "内容生态", "供给"),
    }
    transferable_hits = []
    combined_text = f"{resume_text}\n{job_desc}"
    for group, words in transferable_groups.items():
        resume_hit = any(word in resume_text for word in words)
        job_hit = any(word in job_desc for word in words)
        if resume_hit and job_hit:
            transferable_hits.append(group)
    score = min(92, 45 + keyword_score + len(transferable_hits) * 6)
    missing_words = [word for word in unique_job_words if word not in resume_words][:8]
    mismatch_details = [
        {
            "requirement": f"岗位材料中强调“{word}”相关能力或经验。",
            "resume_evidence": "当前简历没有出现足够直接的对应证据，或只有相邻/泛化描述。",
            "gap": f"缺少能证明“{word}”的具体项目、职责边界、业务场景或结果数据。",
            "impact": "招聘方可能无法判断候选人是否真的做过岗位最看重的工作，而不是只有相邻经验。",
            "action": f"补充一段真实经历：围绕“{word}”说明目标对象、采取动作、协作资源、量化结果和复盘结论。",
        }
        for word in missing_words[:4]
    ]

    return {
        "score": score,
        "summary": "基于岗位关键词、简历证据和可验证成果做的本地估算；配置 API 后可得到更细的顾问式分析。",
        "role_reading": "该岗位需要候选人证明自己能围绕目标内容生态或用户群体，完成供给增长、分层运营、活动策划、数据复盘和跨团队推进。",
        "fit_logic": "本地评分会优先识别创作者/达人运营、生命周期分层、活动策划、数据驱动和协作能力等可迁移证据，再扣除平台、品类、语言或直接经验缺口。",
        "strengths": transferable_hits[:5] or matched[:5] or ["简历已包含部分与岗位相关的运营经历。"],
        "gaps": missing_words[:5] or ["暂无明显关键词缺口。"],
        "suggestions": [
            "优先补充与岗位核心职责一一对应的真实项目，不要只补关键词。",
            "每条补充材料尽量包含场景、动作、协作对象、指标结果和复盘方法。",
            "如果确实没有岗位要求的关键经验，建议降低改写强度，转为强调相邻能力和学习迁移性。",
        ],
        "mismatch_details": mismatch_details,
        "used_ai": False,
        "analysis_error": "未调用 AI，使用本地匹配规则。",
    }


def analyze_resume_match(
    resume_text: str,
    job_desc: str,
    target_role: str = "",
    job_profile: dict | None = None,
) -> dict:
    if not openai.api_key:
        return local_match_analysis(resume_text, job_desc)

    job_profile_text = format_job_profile_for_prompt(job_profile)
    prompt = f"""
你是资深招聘顾问，请评估简历与岗位的匹配度。
你必须先使用“岗位理解”里的业务问题、理想候选人、证据要求、可迁移能力和硬门槛，再评估简历。
不要做关键词重合度判断；要像业务面试官一样判断：候选人的经历能不能证明他解决过类似问题。
重点不是打分，而是解释“岗位真正要什么、简历证明了什么、没证明什么、哪些能力可以迁移、哪些缺口会影响筛选”。

只返回 JSON 对象，不要 Markdown，格式：
{{
  "score": 0-100 的整数,
  "summary": "一句话总结匹配度，要直接说明适配/不适配的核心原因",
  "role_reading": "先用自己的话复述这个岗位真正要解决的业务问题和理想候选人",
  "fit_logic": "说明本次评分逻辑：哪些是核心匹配，哪些只是平台/体裁差异，哪些是硬伤",
  "strengths": ["具体匹配优势1：说明岗位要求和简历证据", "具体匹配优势2", "具体匹配优势3"],
  "gaps": ["具体短板1：岗位要求什么，简历缺什么证据，为什么重要", "具体短板2", "具体短板3"],
  "suggestions": ["具体建议1：补充什么真实材料、放在哪类经历里、最好包含什么指标", "具体建议2", "具体建议3"],
  "mismatch_details": [
    {{
      "requirement": "岗位具体要求，不要只写关键词",
      "resume_evidence": "简历里目前能支撑或不能支撑的证据",
      "gap": "差距是什么，必须具体到场景/经验/能力/指标",
      "impact": "为什么这个差距会影响筛选或面试判断",
      "action": "候选人可以补充什么真实材料，具体到项目、数据、职责、案例"
    }}
  ]
}}

评分标准：
- 经验年限、行业/品类/地区相关性
- 岗位关键词覆盖度和证据强度
- 项目成果、数据复盘、商业化或增长结果
- 核心职责是否做过，而不仅是文字相似
- 表达是否贴近目标岗位的业务场景

评分口径：
1. 先看可迁移能力，不要因为平台名或内容体裁不同就过度降分。比如直播主播/达人运营与 UGC 社区创作者运营，在创作者招募、分层激励、留存活跃、活动策划、数据复盘、跨部门协作上可以高度迁移。
2. 直接经验缺口要扣分，但要区分“硬性门槛缺失”和“平台/内容形态差异”。如果核心运营方法论高度匹配，只是直播 vs 图文/短视频、平台生态不同，分数不应过低。
3. 对有量化成果、完整项目 owner 经历、创作者生命周期管理经验的简历，应给予较高基础分，再指出需要社区化/UGC化表达。
4. 如果岗位明确要求语言、地区、品类、年限等硬条件，而简历没有证据，再进行明显扣分。

诊断要求：
1. 先输出 role_reading 和 fit_logic，证明你已经吃透 JD，而不是直接罗列关键词。
2. strengths 必须引用简历中的具体证据，例如“签约1000+主播”“留存率提升37%”“千万级曝光活动”等；没有证据就不要写。
3. gaps 不要只返回“缺少 XX 经验”这类短句，要说明岗位为什么看重它，以及简历目前缺少哪种证据。
4. mismatch_details 最多 5 条，优先分析会直接影响筛选的硬伤：行业/地区/品类经验、语言要求、年限、核心业务动作、商业化/数据能力。
5. action 必须具体可执行，例如“补充某段欧美用户运营复盘，包含目标用户、活动机制、分层策略、LTV/留存/付费率变化”，不要写“提升匹配度”这类空话。
6. 不要建议候选人编造没有的经历；如果没有相关经验，要建议用相邻经验证明迁移能力，并说明风险。
7. 如果差距只是平台属性或内容体裁差异，要明确说“这是表达和案例补强问题，不是核心能力完全不匹配”。

目标岗位：{target_role or "未指定"}

岗位描述：
{job_desc}

岗位理解：
{job_profile_text}

简历文本：
{resume_text}
""".strip()

    try:
        data = parse_json_object(call_chat_completion(prompt))
        score = int(data.get("score", 0))
        data["score"] = max(0, min(100, score))
        for key in ("strengths", "gaps", "suggestions"):
            value = data.get(key, [])
            if isinstance(value, str):
                data[key] = [value]
            elif not isinstance(value, list):
                data[key] = []
        details = data.get("mismatch_details", [])
        if not isinstance(details, list):
            details = []
        normalized_details = []
        for item in details[:5]:
            if not isinstance(item, dict):
                continue
            normalized_details.append(
                {
                    "requirement": str(item.get("requirement", "")).strip(),
                    "resume_evidence": str(item.get("resume_evidence", "")).strip(),
                    "gap": str(item.get("gap", "")).strip(),
                    "impact": str(item.get("impact", "")).strip(),
                    "action": str(item.get("action", "")).strip(),
                }
            )
        data["mismatch_details"] = [
            item for item in normalized_details if item["requirement"] or item["gap"] or item["action"]
        ]
        if not data["mismatch_details"] and data.get("score", 100) < 70:
            data["mismatch_details"] = local_match_analysis(resume_text, job_desc).get("mismatch_details", [])
        data["summary"] = str(data.get("summary", "")).strip() or "已完成岗位匹配度分析。"
        data["role_reading"] = str(data.get("role_reading", "")).strip()
        data["fit_logic"] = str(data.get("fit_logic", "")).strip()
        data["used_ai"] = True
        data["analysis_error"] = ""
        return data
    except Exception as exc:
        data = local_match_analysis(resume_text, job_desc)
        data["analysis_error"] = f"匹配分析 AI 调用失败，已退回本地规则：{exc}"
        return data


def analyze_role_and_resume_fit(resume_text: str, job_desc: str, target_role: str = "") -> tuple[dict, dict]:
    if not openai.api_key:
        profile = local_job_profile(job_desc, target_role)
        return profile, local_match_analysis(resume_text, job_desc)

    prompt = f"""
你是资深招聘顾问。请一次性完成两件事：
1. 先深读 JD，拆出岗位真正的业务问题、理想候选人、证据要求、可迁移能力和硬门槛。
2. 再基于这份岗位拆解评估简历匹配度。

严禁只做关键词匹配。你必须解释“招聘方为什么要这个人”“简历用什么证据证明能做”“缺口是核心能力缺失还是平台/体裁差异”。

只返回 JSON 对象，不要 Markdown，格式：
{{
  "job_profile": {{
    "role_summary": "一句话说明岗位本质",
    "business_problem": "这个岗位要解决的业务问题",
    "ideal_candidate": "理想候选人画像",
    "core_tasks": ["核心任务"],
    "must_haves": ["硬性能力"],
    "evidence_requirements": ["简历必须出现的证据形态"],
    "transferable_abilities": ["可从相邻经验迁移的能力"],
    "hard_gates": ["不能靠话术弥补的硬门槛"],
    "hidden_criteria": ["隐性筛选点"],
    "resume_strategy": "自然改写策略",
    "keywords": ["关键词"]
  }},
  "match_analysis": {{
    "score": 0-100,
    "summary": "一句话总结匹配度",
    "role_reading": "用自己的话复述岗位真正要什么",
    "fit_logic": "说明评分逻辑，区分核心匹配、可迁移、硬伤、表达差异",
    "strengths": ["引用简历具体证据的优势"],
    "gaps": ["具体短板：岗位要求、简历缺证据、为什么重要"],
    "suggestions": ["具体建议：补什么真实材料，放哪类经历，包含什么指标"],
    "mismatch_details": [
      {{
        "requirement": "岗位具体要求",
        "resume_evidence": "简历当前证据",
        "gap": "具体差距",
        "impact": "筛选影响",
        "action": "具体补救动作"
      }}
    ]
  }}
}}

评分口径：
- 先看岗位要解决的业务问题，再看简历证据，不要反过来找关键词。
- 对直播主播/达人运营到 UGC 社区创作者运营这类场景，要判断创作者生命周期、分层激励、活动策划、数据复盘、跨团队协作是否可迁移。
- 如果只是平台属性或内容体裁差异，要说明这是“表达/案例补强问题”，不要当作核心能力完全不匹配。
- 只有语言、地区、品类、年限、明确海外经验等硬门槛缺失，才明显扣分。
- strengths 必须引用简历证据；suggestions 必须具体到可补材料和指标。

目标岗位：{target_role or "未指定"}

岗位 JD：
{job_desc}

简历：
{resume_text}
""".strip()

    try:
        data = parse_json_object(call_chat_completion(prompt))
        profile = local_job_profile(job_desc, target_role)
        profile_data = data.get("job_profile", {})
        if isinstance(profile_data, dict):
            for key in ("role_summary", "resume_strategy", "business_problem", "ideal_candidate"):
                if profile_data.get(key):
                    profile[key] = str(profile_data[key]).strip()
            for key in (
                "core_tasks",
                "must_haves",
                "hidden_criteria",
                "keywords",
                "evidence_requirements",
                "transferable_abilities",
                "hard_gates",
            ):
                value = profile_data.get(key, [])
                if isinstance(value, str):
                    profile[key] = [value]
                elif isinstance(value, list):
                    profile[key] = [str(item).strip() for item in value if str(item).strip()][:8]
        profile["used_ai"] = True
        profile["analysis_error"] = ""

        match_data = data.get("match_analysis", {})
        if not isinstance(match_data, dict):
            raise ValueError("AI 返回缺少 match_analysis")
        score = int(match_data.get("score", 0))
        match_data["score"] = max(0, min(100, score))
        for key in ("strengths", "gaps", "suggestions"):
            value = match_data.get(key, [])
            if isinstance(value, str):
                match_data[key] = [value]
            elif not isinstance(value, list):
                match_data[key] = []
        details = match_data.get("mismatch_details", [])
        if not isinstance(details, list):
            details = []
        match_data["mismatch_details"] = [
            {
                "requirement": str(item.get("requirement", "")).strip(),
                "resume_evidence": str(item.get("resume_evidence", "")).strip(),
                "gap": str(item.get("gap", "")).strip(),
                "impact": str(item.get("impact", "")).strip(),
                "action": str(item.get("action", "")).strip(),
            }
            for item in details[:5]
            if isinstance(item, dict)
        ]
        match_data["summary"] = str(match_data.get("summary", "")).strip() or "已完成岗位匹配度分析。"
        match_data["role_reading"] = str(match_data.get("role_reading", "")).strip()
        match_data["fit_logic"] = str(match_data.get("fit_logic", "")).strip()
        match_data["used_ai"] = True
        match_data["analysis_error"] = ""
        return profile, match_data
    except Exception as exc:
        profile = local_job_profile(job_desc, target_role)
        profile["analysis_error"] = f"岗位深读+匹配分析 AI 调用失败，已退回本地规则：{exc}"
        match_data = local_match_analysis(resume_text, job_desc)
        match_data["analysis_error"] = f"岗位深读+匹配分析 AI 调用失败，已退回本地规则：{exc}"
        return profile, match_data


def collect_achievement_materials(files, pasted_text: str = "") -> tuple[list[dict], list[str]]:
    materials = []
    errors = []

    if pasted_text.strip():
        materials.append(
            {
                "filename": "手动补充内容",
                "text": pasted_text.strip(),
            }
        )

    for file in files:
        if not file or not file.filename:
            continue

        filename = secure_filename(file.filename) or file.filename
        if not allowed_support_file(file.filename):
            errors.append(f"{file.filename} 暂不支持，已跳过。")
            continue

        data = file.read()
        text, error = extract_text_from_bytes(data, file.filename)
        if error:
            errors.append(error)
            continue

        materials.append({"filename": file.filename, "text": text, "size": len(data)})

    return materials, errors


def collect_achievement_materials_from_payloads(file_payloads: list[dict], pasted_text: str = "") -> tuple[list[dict], list[str]]:
    materials = []
    errors = []

    if pasted_text.strip():
        materials.append({"filename": "手动补充内容", "text": pasted_text.strip()})

    for payload in file_payloads:
        filename = payload.get("filename", "")
        data = payload.get("data", b"")
        if not filename:
            continue
        if not allowed_support_file(filename):
            errors.append(f"{filename} 暂不支持，已跳过。")
            continue
        text, error = extract_text_from_bytes(data, filename)
        if error:
            errors.append(error)
            continue
        materials.append({"filename": filename, "text": text, "size": len(data)})

    return materials, errors


def format_achievement_materials(materials: list[dict], max_chars: int = 12000) -> str:
    chunks = []
    used = 0
    for item in materials:
        filename = item.get("filename", "未命名材料")
        text = str(item.get("text", "")).strip()
        if not text:
            continue
        remaining = max_chars - used
        if remaining <= 0:
            break
        clipped = text[:remaining]
        chunks.append(f"【{filename}】\n{clipped}")
        used += len(clipped)
    return "\n\n".join(chunks)


def local_achievement_analysis(materials: list[dict], job_desc: str) -> dict:
    material_text = format_achievement_materials(materials, max_chars=6000)
    job_words = [
        word
        for word in re.findall(r"[\u4e00-\u9fffA-Za-z0-9+#.]{2,}", job_desc)
        if len(word) >= 2
    ]
    keywords = list(dict.fromkeys(job_words))[:12]
    snippets = []
    for line in material_text.splitlines():
        clean = line.strip()
        if len(clean) < 12:
            continue
        if any(keyword in clean for keyword in keywords):
            snippets.append(clean[:160])
        if len(snippets) >= 5:
            break

    return {
        "summary": "已读取补充材料，并基于岗位关键词做了本地粗筛。",
        "items": [
            {
                "evidence": snippet,
                "relevance": "命中岗位关键词，可作为简历补充素材。",
                "suggested_use": "补充到相关工作经历或项目经历段落中。",
                "confidence": "中",
            }
            for snippet in snippets
        ],
    }


def analyze_achievement_materials(
    materials: list[dict],
    resume_text: str,
    job_desc: str,
    target_role: str = "",
    job_profile: dict | None = None,
) -> dict:
    if not materials:
        return {"summary": "未上传过往成果材料。", "items": []}

    material_text = format_achievement_materials(materials)
    if not material_text:
        return {"summary": "过往成果材料未读取到有效文本。", "items": []}

    if os.getenv("USE_AI_ACHIEVEMENT_ANALYSIS", "").strip() != "1" or not openai.api_key:
        return local_achievement_analysis(materials, job_desc)
    job_profile_text = format_job_profile_for_prompt(job_profile)

    prompt = f"""
你是资深简历顾问。请从“过往工作成果材料”中识别能补充进简历、并且与目标岗位匹配的真实信息。

只返回 JSON 对象，不要 Markdown，格式：
{{
  "summary": "一句话总结补充材料可用性",
  "items": [
    {{
      "evidence": "材料中的原始成果或事实，必须来自材料",
      "relevance": "为什么它匹配岗位",
      "suggested_use": "建议补充到简历的哪个方向/段落",
      "confidence": "高/中/低"
    }}
  ]
}}

要求：
1. 只能提取材料里确实出现的信息，不要编造。
2. 优先筛选与“岗位理解”中的核心任务、硬性能力、隐性筛选点直接相关的成果。
3. 如果材料只有泛泛描述，也可以提炼为“可补充方向”，但 confidence 要标为低。
4. 多份材料要综合判断，去重合并相似成果，不要只看第一份。
5. items 最多返回 10 条。

目标岗位：{target_role or "未指定"}

岗位描述：
{job_desc}

岗位理解：
{job_profile_text}

当前简历：
{resume_text[:8000]}

过往工作成果材料：
{material_text}
""".strip()

    try:
        data = parse_json_object(call_chat_completion(prompt))
        items = data.get("items", [])
        if not isinstance(items, list):
            items = []
        normalized_items = []
        for item in items[:8]:
            if not isinstance(item, dict):
                continue
            normalized_items.append(
                {
                    "evidence": str(item.get("evidence", "")).strip(),
                    "relevance": str(item.get("relevance", "")).strip(),
                    "suggested_use": str(item.get("suggested_use", "")).strip(),
                    "confidence": str(item.get("confidence", "")).strip() or "中",
                }
            )
        return {
            "summary": str(data.get("summary", "")).strip() or "已完成补充材料分析。",
            "items": [item for item in normalized_items if item["evidence"]],
        }
    except Exception:
        return local_achievement_analysis(materials, job_desc)


def run_resume_job(job_id: str, payload: dict) -> None:
    try:
        set_job_progress(job_id, 10, "正在解析简历...")
        original_filename = payload["original_filename"]
        ext = original_filename.rsplit(".", 1)[-1].lower()
        resume_data = payload["resume_data"]
        job_desc = payload["job_desc"]
        target_role = payload["target_role"]
        style = payload["style"]
        format_warning = PDF_FORMAT_WARNING if ext == "pdf" else ""

        original_docx_bytes = None
        original_paragraphs = []
        if ext == "docx":
            original_docx_bytes = resume_data
            original_paragraphs = docx_paragraph_texts(original_docx_bytes)
            resume_text = "\n".join(original_paragraphs).strip()
            if not resume_text:
                raise ValueError("没有从 Word 简历中读取到有效文本。")
        else:
            resume_text, error = extract_text_from_bytes(resume_data, original_filename)
            if error:
                raise ValueError(error)

        set_job_progress(job_id, 22, "正在解析成果材料...")
        achievement_materials, material_errors = collect_achievement_materials_from_payloads(
            payload["achievement_file_payloads"],
            payload["achievement_text"],
        )

        set_job_progress(job_id, 36, "正在理解岗位要求和简历匹配度...")
        job_profile, original_match_analysis = analyze_role_and_resume_fit(resume_text, job_desc, target_role)
        set_job_progress(job_id, 52, "正在筛选可补充的真实成果...")
        achievement_analysis = analyze_achievement_materials(
            achievement_materials,
            resume_text,
            job_desc,
            target_role,
            job_profile,
        )

        set_job_progress(job_id, 68, "正在生成针对岗位的简历改写...")
        optimized_paragraphs = optimize_resume(
            resume_text,
            job_desc,
            target_role,
            style,
            paragraphs=original_paragraphs if original_docx_bytes else None,
            achievement_analysis=achievement_analysis,
            job_profile=job_profile,
        )
        format_protection = {"count": 0, "paragraphs": []}
        if original_docx_bytes:
            raw_optimized_paragraphs = optimized_paragraphs
            optimized_paragraphs = normalize_optimizations_for_docx(original_paragraphs, optimized_paragraphs)
            format_protection = summarize_format_protection(original_paragraphs, raw_optimized_paragraphs, optimized_paragraphs)

        set_job_progress(job_id, 84, "正在整理 Word 文件...")
        OUTPUT_DIR.mkdir(exist_ok=True)
        filename = f"optimized_resume_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_path = OUTPUT_DIR / filename
        if original_docx_bytes:
            output_path.write_bytes(build_optimized_docx_bytes(original_docx_bytes, optimized_paragraphs))
        else:
            output_doc = build_docx_from_text(optimized_paragraphs)
            output_doc.save(output_path)

        set_job_progress(job_id, 92, "正在整理修改说明...")
        change_originals = original_paragraphs if original_docx_bytes else resume_text.splitlines()
        changes = build_change_log(change_originals, optimized_paragraphs)
        display_text = "\n".join(item.get("text", "") for item in optimized_paragraphs)
        if os.getenv("ANALYZE_OPTIMIZED_MATCH", "").strip() == "1":
            optimized_match_analysis = analyze_resume_match(display_text, job_desc, target_role, job_profile)
        else:
            optimized_match_analysis = dict(original_match_analysis)
            optimized_match_analysis["summary"] = "已完成简历生成。为提升线上稳定性，本次跳过优化后二次 AI 匹配评估。"
            optimized_match_analysis["analysis_error"] = "已跳过优化后二次 AI 匹配评估。"
        match_analysis = original_match_analysis
        optimization_diagnosis = build_optimization_diagnosis(
            change_originals,
            optimized_paragraphs,
            changes,
            match_analysis,
            format_protection,
        )
        result_context = {
            "optimized": display_text,
            "download": filename,
            "format_warning": format_warning,
            "job_id": job_id,
            "upload_retention_notice": f"原始上传文件会临时保存 {UPLOAD_TTL_SECONDS // 60} 分钟，可在本页删除。",
            "edit_count": len(changes),
            "changes": changes,
            "match_analysis": match_analysis,
            "optimized_match_analysis": optimized_match_analysis,
            "achievement_analysis": achievement_analysis,
            "job_profile": job_profile,
            "material_count": len(achievement_materials),
            "material_errors": material_errors,
            "optimization_diagnosis": optimization_diagnosis,
            "used_ai": bool(openai.api_key),
        }
        set_job_progress(job_id, 100, "生成完成，可以查看结果并下载 Word。", status="done", result=result_context)
    except Exception as exc:
        traceback.print_exc()
        set_job_progress(job_id, 100, f"生成失败：{exc}", status="failed", error=str(exc))


@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "GET":
        return render_index()

    resume_file = request.files.get("resume")
    job_desc = request.form.get("job_desc", "").strip()
    job_file = request.files.get("job_file")
    target_role = request.form.get("target_role", "").strip()
    style = request.form.get("style", "professional")
    achievement_text = request.form.get("achievement_text", "").strip()

    if password_enabled() and not is_authenticated():
        password = request.form.get("password", "")
        if not secrets.compare_digest(password, APP_PASSWORD):
            return render_index(error="生成优化前请输入正确口令。", form=request.form), 401
        session["authenticated"] = True

    if not resume_file or not resume_file.filename:
        return render_index(error="请上传一份 PDF 或 DOCX 简历。")
    if not allowed_file(resume_file.filename):
        return render_index(error="文件格式不支持，请上传 PDF 或 DOCX。")

    job_file_error = None
    if job_file and job_file.filename:
        file_job_desc, job_file_error = extract_job_description_from_file(job_file)
        if file_job_desc:
            job_desc = "\n\n".join(part for part in (job_desc, file_job_desc) if part)
    if not job_desc:
        return render_index(error=job_file_error or "请填写岗位描述，或上传岗位页面/JD 文件。", form=request.form)

    original_filename = secure_filename(resume_file.filename)
    resume_data = resume_file.read()
    job_id = uuid.uuid4().hex
    saved_uploads = [save_upload_file(job_id, "resume", original_filename, resume_data)]
    achievement_file_payloads = [
        {"filename": file.filename, "data": file.read()}
        for file in request.files.getlist("achievement_files")
        if file and file.filename
    ]
    for index, payload_item in enumerate(achievement_file_payloads, start=1):
        saved_uploads.append(save_upload_file(job_id, f"achievement_{index}", payload_item["filename"], payload_item["data"]))
    with JOBS_LOCK:
        JOBS[job_id] = {
            "status": "running",
            "percent": 3,
            "message": "任务已创建，正在准备生成...",
            "saved_uploads": saved_uploads,
            "created_at": datetime.now().isoformat(timespec="seconds"),
            "updated_at": datetime.now().isoformat(timespec="seconds"),
        }
    payload = {
        "original_filename": original_filename,
        "resume_data": resume_data,
        "job_desc": job_desc,
        "target_role": target_role,
        "style": style,
        "achievement_text": achievement_text,
        "achievement_file_payloads": achievement_file_payloads,
    }
    if original_filename.rsplit(".", 1)[-1].lower() == "pdf":
        set_job_progress(job_id, 4, PDF_FORMAT_WARNING)
    worker = threading.Thread(target=run_resume_job, args=(job_id, payload), daemon=True)
    worker.start()
    return redirect(url_for("job_status_page", job_id=job_id))


@app.get("/jobs/<job_id>")
def job_status_page(job_id: str):
    if not get_job(job_id):
        return render_template("error.html", message="任务不存在或已过期，请重新提交。"), 404
    return render_template("job.html", job_id=job_id, upload_ttl_minutes=UPLOAD_TTL_SECONDS // 60)


@app.get("/jobs/<job_id>/status")
def job_status(job_id: str):
    job = get_job(job_id)
    if not job:
        return {"status": "missing", "percent": 100, "message": "任务不存在或已过期，请重新提交。"}, 404
    return {
        "status": job.get("status", "running"),
        "percent": job.get("percent", 0),
        "message": job.get("message", ""),
        "result_url": url_for("job_result", job_id=job_id) if job.get("status") == "done" else "",
        "delete_uploads_url": url_for("delete_uploads", job_id=job_id),
        "uploads_saved": bool(job.get("saved_uploads")),
        "error": job.get("error", ""),
    }


@app.get("/jobs/<job_id>/result")
def job_result(job_id: str):
    job = get_job(job_id)
    if not job:
        return render_template("error.html", message="任务不存在或已过期，请重新提交。"), 404
    if job.get("status") == "failed":
        return render_template("error.html", message=job.get("error") or "生成失败，请重新提交。"), 500
    if job.get("status") != "done":
        return redirect(url_for("job_status_page", job_id=job_id))
    return render_template("result.html", **job.get("result", {}))


@app.post("/jobs/<job_id>/uploads/delete")
def delete_uploads(job_id: str):
    shutil.rmtree(UPLOAD_DIR / secure_filename(job_id), ignore_errors=True)
    with JOBS_LOCK:
        if job_id in JOBS:
            JOBS[job_id]["saved_uploads"] = []
            JOBS[job_id]["uploads_deleted"] = True
    next_url = request.form.get("next") or url_for("job_status_page", job_id=job_id)
    return redirect(next_url)


@app.get("/download/<path:filename>")
def download_file(filename: str):
    safe_name = secure_filename(filename)
    if safe_name != filename or not safe_name.startswith("optimized_resume_") or not safe_name.endswith(".docx"):
        return render_template("error.html", message="下载文件不存在或文件名无效。"), 404
    return send_from_directory(
        OUTPUT_DIR,
        safe_name,
        as_attachment=True,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        download_name=safe_name,
    )


@app.route("/admin/login", methods=["GET", "POST"])
def admin_login():
    if not admin_enabled():
        return render_template("error.html", message="管理功能未启用。请先在 Railway 设置 ADMIN_PASSWORD。"), 404

    error = ""
    if request.method == "POST":
        password = request.form.get("password", "")
        if secrets.compare_digest(password, ADMIN_PASSWORD):
            session["admin_authenticated"] = True
            return redirect(request.args.get("next") or url_for("admin_uploads"))
        error = "管理员口令不正确，请重试。"
    return render_template("login.html", error=error, title="管理员口令", description="请输入管理员口令查看临时上传文件。")


@app.post("/admin/logout")
def admin_logout():
    session.pop("admin_authenticated", None)
    return redirect(url_for("admin_login"))


@app.get("/admin/uploads")
def admin_uploads():
    auth_response = require_admin()
    if auth_response:
        return auth_response
    return render_template(
        "admin_uploads.html",
        records=upload_records(),
        ttl_minutes=UPLOAD_TTL_SECONDS // 60,
    )


@app.get("/admin/uploads/<job_id>/<path:filename>")
def admin_download_upload(job_id: str, filename: str):
    auth_response = require_admin()
    if auth_response:
        return auth_response
    safe_job_id = secure_filename(job_id)
    safe_name = secure_filename(filename)
    if safe_job_id != job_id or safe_name != filename:
        return render_template("error.html", message="文件名无效。"), 404
    return send_from_directory(UPLOAD_DIR / safe_job_id, safe_name, as_attachment=True, download_name=safe_name)


@app.errorhandler(413)
def file_too_large(_error):
    return render_index(error="文件太大，请上传 8MB 以内的简历。"), 413


@app.get("/health")
def health():
    return {
        "ok": True,
        "version": APP_VERSION,
        "password_enabled": password_enabled(),
        "deepseek_key_configured": bool(openai.api_key),
        "deepseek_base": openai.api_base,
        "output_dir_exists": OUTPUT_DIR.exists(),
    }


@app.errorhandler(Exception)
def internal_error(error):
    if isinstance(error, HTTPException):
        return error
    traceback.print_exc()
    message = "服务器处理失败，请稍后重试。"
    if os.getenv("SHOW_ERROR_DETAILS", "").strip() == "1":
        message = f"{type(error).__name__}: {error}"
    return render_template("error.html", message=message), 500


if __name__ == "__main__":
    app.run(host=os.getenv("HOST", "127.0.0.1"), port=int(os.getenv("PORT", "5000")), debug=False, use_reloader=False)
